import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def set_cell_border(cell, **kwargs):
    """
    kwargs can be top, bottom, left, right
    values: {"sz": 12, "val": "single", "color": "1E3A8A"}
    """
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            b_xml = parse_xml(f'<w:{edge} {nsdecls("w")} w:val="{edge_data.get("val", "single")}" w:sz="{edge_data.get("sz", 4)}" w:space="0" w:color="{edge_data.get("color", "auto")}"/>')
            tcBorders.append(b_xml)
        else:
            b_xml = parse_xml(f'<w:{edge} {nsdecls("w")} w:val="none"/>')
            tcBorders.append(b_xml)
    tcPr.append(tcBorders)

def add_callout(doc, text_list, title="CATATAN PENTING", callout_type="info"):
    # colors
    if callout_type == "warning":
        bg_color = "FEF3C7"  # Amber 100
        border_color = "D97706"  # Amber 600
        title_color = RGBColor(180, 83, 9)
    elif callout_type == "success":
        bg_color = "DCFCE7"  # Green 100
        border_color = "16A34A"  # Green 600
        title_color = RGBColor(22, 101, 52)
    elif callout_type == "formula":
        bg_color = "F0FDF4"  # Green 50
        border_color = "059669"  # Emerald 600
        title_color = RGBColor(4, 120, 87)
    else: # info
        bg_color = "EFF6FF"  # Blue 50
        border_color = "2563EB"  # Blue 600
        title_color = RGBColor(29, 78, 216)

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, bg_color)
    set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
    set_cell_border(cell, left={"sz": 24, "val": "single", "color": border_color},
                          top={"val": "none"}, bottom={"val": "none"}, right={"val": "none"})
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(3)
    r_title = p.add_run(f"📌 {title}\n" if title else "")
    r_title.bold = True
    r_title.font.name = "Arial"
    r_title.font.size = Pt(10.5)
    r_title.font.color.rgb = title_color
    
    for idx, item in enumerate(text_list):
        if idx > 0 or title:
            p_item = cell.add_paragraph()
            p_item.paragraph_format.space_before = Pt(2)
            p_item.paragraph_format.space_after = Pt(2)
        else:
            p_item = p
        
        r_item = p_item.add_run(item)
        r_item.font.name = "Arial"
        r_item.font.size = Pt(9.5)
        r_item.font.color.rgb = RGBColor(31, 41, 55)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def build_word_guide():
    doc = Document()
    
    # Page setup - Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

    # Style defaults
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Arial'
    style_normal.font.size = Pt(10)
    style_normal.font.color.rgb = RGBColor(31, 41, 55)
    style_normal.paragraph_format.line_spacing = 1.15
    style_normal.paragraph_format.space_after = Pt(4)

    # --- COVER / TITLE BLOCK ---
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(12)
    title_p.paragraph_format.space_after = Pt(2)
    r_badge = title_p.add_run("MODUL MONITORING & MANAJEMEN BIAYA")
    r_badge.bold = True
    r_badge.font.size = Pt(9)
    r_badge.font.color.rgb = RGBColor(16, 185, 129)

    h_main = doc.add_paragraph()
    h_main.paragraph_format.space_before = Pt(2)
    h_main.paragraph_format.space_after = Pt(4)
    r_main = h_main.add_run("BUKU PANDUAN PENGGUNAAN APLIKASI\nKMT CORN (DIVISI JAGUNG)")
    r_main.bold = True
    r_main.font.size = Pt(18)
    r_main.font.color.rgb = RGBColor(30, 58, 138) # Navy

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_before = Pt(0)
    sub_p.paragraph_format.space_after = Pt(12)
    r_sub = sub_p.add_run("Sistem Informasi ERP PT Karisma Indoagro Universal\nMonitoring Biaya Operasional, DCA, Omset, & Analisis Cost / Hasil")
    r_sub.italic = True
    r_sub.font.size = Pt(10.5)
    r_sub.font.color.rgb = RGBColor(75, 85, 99)

    # Divider bar
    div_table = doc.add_table(rows=1, cols=1)
    div_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    div_cell = div_table.cell(0, 0)
    div_cell.width = Inches(6.5)
    set_cell_background(div_cell, "1E3A8A")
    set_cell_margins(div_cell, top=10, bottom=10, left=0, right=0)
    p_div = div_cell.paragraphs[0]
    p_div.paragraph_format.space_before = Pt(0)
    p_div.paragraph_format.space_after = Pt(0)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # --- DAFTAR ISI TABLE ---
    toc_head = doc.add_heading(level=2)
    r_toc = toc_head.add_run("DAFTAR ISI")
    r_toc.font.size = Pt(13)
    r_toc.font.color.rgb = RGBColor(30, 58, 138)
    r_toc.bold = True
    toc_head.paragraph_format.space_before = Pt(8)
    toc_head.paragraph_format.space_after = Pt(6)

    toc_items = [
        ("1. Tentang Modul KMT CORN", "Gambaran umum, tujuan sistem, dan indikator efisiensi Cost/Hasil"),
        ("2. Hak Akses & Peran Pengguna (User Roles)", "Matriks wewenang Super Admin / KADEP, Admin Keuangan, & ABM"),
        ("3. Dashboard KMT CORN (Cost / Hasil YTD)", "Filter data, KPI Cards, Rekapitulasi Bulanan YTD, dan Export Excel"),
        ("4. Modul Data Omset & Retur Penjualan", "Pencatatan faktur jagung, produk BISI 959 / Q-235, & kalkulasi retur"),
        ("5. Modul Biaya DCA (Demonstration & Cost Activities)", "Formulir multi-kegiatan, Uang Muka, Realisasi, Refund & Rekap Hierarkis"),
        ("6. Modul Biaya Operasional Lapangan", "15 pos beban operasional sales/MDO/ABM, uang muka, & refund"),
        ("7. Modul Promo Material & Peralatan", "Pencatatan spanduk, brosur, alat uji kadar air, timbangan, & rak"),
        ("8. Modul Biaya Lain-Lain (Others)", "Pencatatan perizinan, donasi CSR demo plot, & perbaikan sarana"),
        ("9. Modul Penggajian (Gaji Karyawan Lapangan)", "Matriks gaji 12 bulan karyawan lapangan & integrasi biaya"),
        ("10. Workflow Verifikasi & Kontrol Keuangan", "Alur review, status verifikasi, penguncian data, & audit reset otomatis"),
        ("11. Panduan Import & Export Excel", "Langkah import data massal template Excel & opsi export laporan"),
        ("12. Pertanyaan Umum (FAQ) & Solusi Masalah", "Jawaban atas kendala operasional yang sering dihadapi pengguna")
    ]

    toc_table = doc.add_table(rows=len(toc_items) + 1, cols=2)
    toc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    toc_table.autofit = False

    # Header TOC
    th1 = toc_table.cell(0, 0)
    th2 = toc_table.cell(0, 1)
    th1.width = Inches(2.8)
    th2.width = Inches(3.7)
    set_cell_background(th1, "F3F4F6")
    set_cell_background(th2, "F3F4F6")
    set_cell_margins(th1, top=80, bottom=80, left=100, right=100)
    set_cell_margins(th2, top=80, bottom=80, left=100, right=100)
    set_cell_border(th1, bottom={"sz": 8, "val": "single", "color": "9CA3AF"})
    set_cell_border(th2, bottom={"sz": 8, "val": "single", "color": "9CA3AF"})

    r_th1 = th1.paragraphs[0].add_run("BAB / TOPIK")
    r_th1.bold = True
    r_th1.font.size = Pt(9)
    r_th1.font.color.rgb = RGBColor(55, 65, 81)
    r_th2 = th2.paragraphs[0].add_run("DESKRIPSI CAKUPAN")
    r_th2.bold = True
    r_th2.font.size = Pt(9)
    r_th2.font.color.rgb = RGBColor(55, 65, 81)

    for i, (title_toc, desc_toc) in enumerate(toc_items):
        row_idx = i + 1
        c1 = toc_table.cell(row_idx, 0)
        c2 = toc_table.cell(row_idx, 1)
        c1.width = Inches(2.8)
        c2.width = Inches(3.7)
        if i % 2 == 1:
            set_cell_background(c1, "FAFAFA")
            set_cell_background(c2, "FAFAFA")
        set_cell_margins(c1, top=60, bottom=60, left=100, right=100)
        set_cell_margins(c2, top=60, bottom=60, left=100, right=100)
        set_cell_border(c1, bottom={"sz": 4, "val": "single", "color": "E5E7EB"})
        set_cell_border(c2, bottom={"sz": 4, "val": "single", "color": "E5E7EB"})

        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_before = Pt(1)
        p1.paragraph_format.space_after = Pt(1)
        r1 = p1.add_run(title_toc)
        r1.bold = True
        r1.font.size = Pt(9)
        r1.font.color.rgb = RGBColor(30, 58, 138)

        p2 = c2.paragraphs[0]
        p2.paragraph_format.space_before = Pt(1)
        p2.paragraph_format.space_after = Pt(1)
        r2 = p2.add_run(desc_toc)
        r2.font.size = Pt(8.5)
        r2.font.color.rgb = RGBColor(75, 85, 99)

    doc.add_page_break()

    # --- HELPER FUNCTION FOR SECTION HEADERS ---
    def add_section_header(num_title, subtitle=""):
        h = doc.add_heading(level=1)
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(4)
        r = h.add_run(num_title)
        r.bold = True
        r.font.size = Pt(13)
        r.font.color.rgb = RGBColor(30, 58, 138)
        
        if subtitle:
            p_sub = doc.add_paragraph()
            p_sub.paragraph_format.space_before = Pt(0)
            p_sub.paragraph_format.space_after = Pt(6)
            r_sub = p_sub.add_run(f"📍 Menu / URL: {subtitle}")
            r_sub.font.size = Pt(9)
            r_sub.font.color.rgb = RGBColor(5, 150, 105)
            r_sub.bold = True

    # --- BAB 1 ---
    add_section_header("1. Tentang Modul KMT CORN")
    
    p = doc.add_paragraph()
    p.add_run("Modul ").font.size = Pt(9.5)
    r_bold = p.add_run("KMT CORN (Komoditas Jagung)")
    r_bold.bold = True
    p.add_run(" dirancang khusus untuk memonitoring, mencatat, mengontrol, dan menganalisis performa bisnis divisi jagung PT Karisma Indoagro Universal (seperti varietas unggul benih ")
    p.add_run("BISI 959, Q-235 CLING").bold = True
    p.add_run(", dan varietas komersial lainnya).\n\nTujuan utama sistem ini adalah menyediakan visibilitas finansial yang akurat dan transparan dengan menghitung rasio efisiensi biaya terhadap hasil penjualan (")
    p.add_run("Rasio Cost / Hasil").bold = True
    p.add_run(") secara real-time, baik secara konsolidasi nasional, per area/wilayah pemasaran, maupun tren bulanan (Year to Date / YTD).")

    add_callout(doc, [
        "Rumus Perhitungan:",
        "Cost / Hasil (%) = (Total Seluruh Biaya / Total Omset Penjualan Neto) x 100%",
        "",
        "Kategori Indikator Efisiensi:",
        "• Hijau (< 20.00%)       : Efisien / Sangat Aman (Target operasional optimal tercapai)",
        "• Kuning (20.00% - 30.00%) : Waspada / Perhatian (Biaya mendekati batas ambang toleransi)",
        "• Merah (> 30.00%)        : Kritis / Over Budget (Biaya melebihi batas efisiensi, evaluasi diperlukan)"
    ], title="FORMULA & INDIKATOR KUNCI COST / HASIL", callout_type="formula")

    # --- BAB 2 ---
    add_section_header("2. Hak Akses & Peran Pengguna (User Roles)")
    
    p = doc.add_paragraph()
    p.add_run("Untuk menjaga tata kelola dan pemisahan fungsi (*segregation of duties*), sistem membagi kewenangan pengguna ke dalam 3 tingkatan peran:")

    role_table_data = [
        ("Fitur / Menu", "Level 1: Super Admin / KADEP", "Level 2: ADMKEU (Admin Finance)", "Level 3: ABM (Area Manager)"),
        ("Cakupan Wilayah", "Seluruh Wilayah (Nasional)", "Seluruh Wilayah (Nasional)", "Hanya Wilayah Tugas Sendiri"),
        ("Dashboard Cost / Hasil", "Full Access + Export Excel", "Full Access + Export Excel", "Read-Only (Wilayah Sendiri)"),
        ("Data Omset & Retur", "Input, Edit, Import, Export", "Input, Edit, Import, Export", "Tidak Dapat Menginput"),
        ("Biaya Operasional", "View, Edit, Verifikasi, Export", "View, Verifikasi, Export", "Input & Edit (Wilayah Sendiri)"),
        ("Biaya DCA Kegiatan", "Input, Edit, Verifikasi, Export", "Verifikasi, Export, Tambah Master", "Input & Edit (Wilayah Sendiri)"),
        ("Promo Material & Alat", "Input, Edit, Import, Export", "Input, Edit, Import, Export", "Tidak Ada Akses"),
        ("Biaya Gaji Lapangan", "Input, Edit, Import, Export", "Tidak Ada Akses (Confidential)", "Tidak Ada Akses"),
        ("Biaya Lain-lain (Others)", "Input, Edit, Export", "Input, Edit, Export", "Tidak Ada Akses")
    ]

    r_tbl = doc.add_table(rows=len(role_table_data), cols=4)
    r_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    r_tbl.autofit = False

    col_widths = [Inches(1.8), Inches(1.6), Inches(1.6), Inches(1.5)]

    for row_idx, row_vals in enumerate(role_table_data):
        for col_idx, text in enumerate(row_vals):
            c = r_tbl.cell(row_idx, col_idx)
            c.width = col_widths[col_idx]
            set_cell_margins(c, top=70, bottom=70, left=80, right=80)
            
            p_cell = c.paragraphs[0]
            p_cell.paragraph_format.space_before = Pt(1)
            p_cell.paragraph_format.space_after = Pt(1)

            if row_idx == 0:
                set_cell_background(c, "1E3A8A")
                set_cell_border(c, bottom={"sz": 12, "val": "single", "color": "1E40AF"},
                                   top={"sz": 4, "val": "single", "color": "1E40AF"},
                                   left={"sz": 4, "val": "single", "color": "1E40AF"},
                                   right={"sz": 4, "val": "single", "color": "1E40AF"})
                r_txt = p_cell.add_run(text)
                r_txt.bold = True
                r_txt.font.size = Pt(8.5)
                r_txt.font.color.rgb = RGBColor(255, 255, 255)
            else:
                bg = "F9FAFB" if row_idx % 2 == 1 else "FFFFFF"
                set_cell_background(c, bg)
                set_cell_border(c, bottom={"sz": 4, "val": "single", "color": "E5E7EB"},
                                   left={"sz": 4, "val": "single", "color": "E5E7EB"},
                                   right={"sz": 4, "val": "single", "color": "E5E7EB"},
                                   top={"sz": 4, "val": "single", "color": "E5E7EB"})
                r_txt = p_cell.add_run(text)
                r_txt.font.size = Pt(8)
                if col_idx == 0:
                    r_txt.bold = True
                    r_txt.font.color.rgb = RGBColor(30, 58, 138)
                else:
                    r_txt.font.color.rgb = RGBColor(55, 65, 81)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    add_callout(doc, [
        "Data yang SUDAH DIVERIFIKASI oleh Admin Keuangan / KADEP terkunci secara otomatis.",
        "Pengguna ABM tidak dapat mengubah atau menghapus data berstatus 'Terverifikasi' demi menjamin keabsahan audit pembukuan kas."
    ], title="ATURAN INTEGRITAS DATA & PENGUNCIAN", callout_type="warning")

    # --- BAB 3 ---
    add_section_header("3. Dashboard KMT CORN (Cost / Hasil YTD)", "kmt/dashboard")
    
    p = doc.add_paragraph()
    p.add_run("Dashboard KMT CORN merupakan pusat kendali eksekutif untuk melihat performa biaya terhadap penjualan secara makro maupun mikro.")

    doc.add_paragraph("Fitur dan Komponen Utama Dashboard:").bold = True
    
    dash_points = [
        ("Filter Periode & Wilayah:", "Pengguna dapat memilih Tahun Buku, rentang Bulan Awal s.d. Bulan Akhir (misal: Semester 1 Jan-Jun atau Q1 Jan-Mar), dan memilih Wilayah spesifik atau Konsolidasi Nasional."),
        ("Kartu Ringkasan KPI (Summary Cards):", "Menampilkan 4 metrik utama: Total Omset (Neto Penjualan Jagung), Total Biaya (Akumulasi Operasional, DCA, Promo, Peralatan, Others, Gaji), Total Gaji SDM Lapangan, dan Rasio Cost/Hasil YTD (berwarna otomatis sesuai ambang batas)."),
        ("Tabel Rekapitulasi YTD Bulanan:", "Menampilkan matriks 12 baris bulan (Januari s.d. Desember) dengan rincian kolom Omset, Operasional, DCA, Peralatan, Others, Gaji, Total Biaya, dan persentase Cost/Hasil tiap bulan."),
        ("Tabel Evaluasi Kuartal per Wilayah:", "Menganalisis pergerakan persentase biaya per kuartal (Q1, Q2, Q3, Q4) dan rata-rata tahunan untuk masing-masing area pemasaran."),
        ("Fitur Export Excel Dashboard:", "Menghasilkan file Excel multi-sheet otomatis (Sheet Konsolidasi Nasional + Sheet per masing-masing Wilayah) dengan formula perhitungan siap cetak.")
    ]

    for title_dp, desc_dp in dash_points:
        p_bp = doc.add_paragraph(style='List Bullet')
        p_bp.paragraph_format.space_before = Pt(1)
        p_bp.paragraph_format.space_after = Pt(2)
        r1 = p_bp.add_run(title_dp + " ")
        r1.bold = True
        r1.font.size = Pt(9.5)
        r2 = p_bp.add_run(desc_dp)
        r2.font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # --- BAB 4 ---
    add_section_header("4. Modul Data Omset & Retur Penjualan", "kmt/omset")

    doc.add_paragraph("4.1. Alur Input Data Omset Penjualan").bold = True
    omset_steps = [
        "Buka menu Omset -> klik tombol Tambah Data di pojok kanan atas tabel.",
        "Pilih Tanggal Penjualan dan Wilayah Pemasaran kios/toko penerima.",
        "Masukkan Nama Toko / Kios, Kota/Kabupaten, dan Kode SC (Sales Code).",
        "Pilih Produk Jagung (misal: BISI 959 atau Q-235 CLING).",
        "Masukkan Quantity (Unit/Sak/Karton) dan Harga Satuan Inc PPN.",
        "Sistem akan otomatis menghitung DPP (Dasar Pengenaan Pajak) dan Total Nilai Penjualan Neto Inc PPN.",
        "Klik Simpan. Data langsung terakumulasi ke dalam Dashboard Cost / Hasil."
    ]
    for i, step in enumerate(omset_steps):
        p_num = doc.add_paragraph(style='List Number')
        p_num.paragraph_format.space_before = Pt(1)
        p_num.paragraph_format.space_after = Pt(2)
        r = p_num.add_run(step)
        r.font.size = Pt(9.5)

    doc.add_paragraph("4.2. Pengelolaan Retur Penjualan Terintegrasi").bold = True
    p = doc.add_paragraph()
    p.add_run("Jika terdapat barang retur dari kios (akibat expired, reject fisik, atau penarikan konsinyasi), catat retur melalui langkah berikut:")
    
    retur_steps = [
        "Pada baris transaksi toko yang bersangkutan di tabel Omset, klik tombol Retur.",
        "Klik Tambah Retur pada jendela modal popup yang muncul.",
        "Isi Nomor Dokumen Retur, Tanggal Retur, Jumlah Qty Retur, dan Nilai Nominal Retur.",
        "Pilih Opsi Kurangi Target Omset:",
        "  • Ya (Kurangi Target)    : Omset neto toko tersebut dipotong otomatis, sehingga angka Omset di Dashboard terkoreksi turun.",
        "  • Tidak (Hanya Catatan) : Retur dicatat murni untuk arsip logistik tanpa mengubah pencapaian nominal omset.",
        "Klik Simpan Retur."
    ]
    for step in retur_steps:
        p_step = doc.add_paragraph()
        p_step.paragraph_format.left_indent = Inches(0.25)
        p_step.paragraph_format.space_before = Pt(1)
        p_step.paragraph_format.space_after = Pt(1)
        r = p_step.add_run(step)
        r.font.size = Pt(9)
        if "Ya (" in step or "Tidak (" in step:
            r.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # --- BAB 5 ---
    add_section_header("5. Modul Biaya DCA (Demonstration & Cost Activities)", "kmt/dca")

    p = doc.add_paragraph()
    p.add_run("Modul DCA mencatat pendanaan dan realisasi kegiatan promosi langsung di lapangan (Farmer Meeting, Big Farmer Meeting / BFM, Farmers Field Day / FFD, ODP, Expo Jagung, dan Demo Plot).")

    doc.add_paragraph("5.1. Pengisian Formulir DCA Multi-Kegiatan").bold = True
    dca_guide = [
        ("Header Dokumen:", "Isi Tanggal DCA, pilih Wilayah & ABM penanggung jawab, tentukan Nama MDO pelaksana, serta masukkan nominal Uang Muka (UM Kasbon) yang diterima dari kasir finance."),
        ("Tabel Dinamis Kegiatan:", "Klik tombol '+ Tambah Kegiatan' untuk menambahkan rincian acara (bisa multi kegiatan dalam 1 nomor kasbon):"),
        ("  • Jenis Kegiatan", "Pilih tipe acara (BFM, FM, FFD, Demo Plot, dll). Bisa tambah master baru langsung."),
        ("  • Tanggal", "Tentukan Tanggal Kasbon dan Tanggal Pelaksanaan Kegiatan."),
        ("  • Peserta & Qty", "Isi estimasi Jumlah Peserta (petani/kios) dan Penjualan Langsung produk BISI 959 / Q-235 saat acara."),
        ("  • Realisasi Biaya", "Masukkan total biaya riil konsumsi, tenda, sound system, dan perlengkapan."),
        ("  • Keterangan", "Tuliskan nama desa, kecamatan, atau kios lokasi kegiatan."),
        ("Kalkulasi Otomatis:", "Sistem otomatis menjumlahkan Total Realisasi = Σ Realisasi Biaya. Jika UM > Total Realisasi, sistem otomatis mengkalkulasi nilai Refund (Sisa Kasbon yang harus disetor kembali ke kas).")
    ]
    for item_h, item_d in dca_guide:
        p_dca = doc.add_paragraph()
        p_dca.paragraph_format.left_indent = Inches(0.2)
        p_dca.paragraph_format.space_before = Pt(1)
        p_dca.paragraph_format.space_after = Pt(1)
        p_dca.add_run(item_h + " ").bold = True
        p_dca.add_run(item_d)
        p_dca.runs[0].font.size = Pt(9.5)
        p_dca.runs[1].font.size = Pt(9.5)

    doc.add_paragraph("5.2. Rekapitulasi Berjenjang & Cetak Laporan").bold = True
    p = doc.add_paragraph()
    p.add_run("• ")
    p.add_run("Rekap Hierarki: ").bold = True
    p.add_run("Melihat data pengeluaran DCA yang dikelompokkan bertingkat: ABM -> MDO -> Jenis Kegiatan.\n")
    p.add_run("• ")
    p.add_run("Export Excel Rincian: ").bold = True
    p.add_run("Menghasilkan laporan resmi format landscape lengkap dengan tanda tangan ABM, Admin Keuangan, dan KADEP.")

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # --- BAB 6 ---
    add_section_header("6. Modul Biaya Operasional Lapangan", "kmt/operasional")

    p = doc.add_paragraph()
    p.add_run("Modul Operasional digunakan untuk mencatat pengeluaran rutin harian tenaga lapangan (Sales, MDO, ABM) saat menjalankan dinas luar kota.")

    doc.add_paragraph("15 Pos Alokasi Biaya Standar:").bold = True

    pos_list = [
        ("1. Hotel / Penginapan", "Biaya sewa kamar hotel dinas luar kota."),
        ("2. Per Diem", "Uang makan dan saku harian perjalanan dinas."),
        ("3. Entertainment", "Biaya jamuan relasi bisnis, kios, atau tamu kedinasan."),
        ("4. Communication", "Paket data internet dan pulsa koordinasi lapangan."),
        ("5. ATK", "Alat tulis kantor, nota, map, dan perlengkapan admin area."),
        ("6. Gasoline (BBM)", "Bahan bakar kendaraan operasional dinas."),
        ("7. Sparepart & Service", "Perawatan rutin dan perbaikan kendaraan dinas."),
        ("8. Tol, Retribusi, Parkir", "Karcis jalan tol, parkir pasar/kios, dan retribusi jalan."),
        ("9. Transportasi", "Tiket bus, travel, kereta api, atau sewa perahu/ojek."),
        ("10. Pos & Paket", "Ongkos kirim sampel benih, faktur, dan berkas ke pusat."),
        ("11. Tambah Angin", "Biaya pengisian angin ban kendaraan operasional."),
        ("12. Tambal Ban", "Biaya perbaikan ban bocor saat bertugas."),
        ("13. Indekost", "Sewa rumah singgah / mess tim lapangan."),
        ("14. Sewa Kendaraan", "Rental mobil/motor operasional kegiatan area."),
        ("15. Lain-lain", "Pengeluaran tak terduga yang dilengkapi bukti sah.")
    ]

    pos_table = doc.add_table(rows=8, cols=2)
    pos_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    pos_table.autofit = False

    for r_idx in range(8):
        # col 0
        idx1 = r_idx
        idx2 = r_idx + 8 if r_idx + 8 < len(pos_list) else None

        c1 = pos_table.cell(r_idx, 0)
        c2 = pos_table.cell(r_idx, 1)
        c1.width = Inches(3.25)
        c2.width = Inches(3.25)
        set_cell_margins(c1, top=40, bottom=40, left=60, right=60)
        set_cell_margins(c2, top=40, bottom=40, left=60, right=60)

        set_cell_border(c1, bottom={"sz": 4, "val": "single", "color": "E5E7EB"},
                            top={"sz": 4, "val": "single", "color": "E5E7EB"},
                            left={"sz": 4, "val": "single", "color": "E5E7EB"},
                            right={"sz": 4, "val": "single", "color": "E5E7EB"})
        set_cell_border(c2, bottom={"sz": 4, "val": "single", "color": "E5E7EB"},
                            top={"sz": 4, "val": "single", "color": "E5E7EB"},
                            left={"sz": 4, "val": "single", "color": "E5E7EB"},
                            right={"sz": 4, "val": "single", "color": "E5E7EB"})

        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_before = Pt(1)
        p1.paragraph_format.space_after = Pt(1)
        r1_b = p1.add_run(pos_list[idx1][0] + ": ")
        r1_b.bold = True
        r1_b.font.size = Pt(8.5)
        r1_d = p1.add_run(pos_list[idx1][1])
        r1_d.font.size = Pt(8)

        p2 = c2.paragraphs[0]
        p2.paragraph_format.space_before = Pt(1)
        p2.paragraph_format.space_after = Pt(1)
        if idx2 is not None:
            r2_b = p2.add_run(pos_list[idx2][0] + ": ")
            r2_b.bold = True
            r2_b.font.size = Pt(8.5)
            r2_d = p2.add_run(pos_list[idx2][1])
            r2_d.font.size = Pt(8)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # --- BAB 7, 8, 9 ---
    add_section_header("7. Modul Promo Material & Peralatan", "kmt/promo")
    p = doc.add_paragraph()
    p.add_run("Mencatat belanja sarana promosi dan inventaris pendukung divisi jagung:\n")
    p.add_run("• Promo Material: ").bold = True
    p.add_run("Spanduk kios, banner jalan, kaos promosi petani, topi, brosur varietas.\n")
    p.add_run("• Peralatan Lapangan: ").bold = True
    p.add_run("Timbangan digital jagung, alat ukur kadar air (moisture tester), tenda display, dan rak pajang kios.\nSetiap transaksi mencatat tanggal, vendor/supplier, nama item, nilai biaya, dan wilayah penempatan.")

    add_section_header("8. Modul Biaya Lain-Lain (Others)", "kmt/others")
    p = doc.add_paragraph()
    p.add_run("Mencatat pengeluaran yang tidak termasuk dalam pos DCA, Operasional Harian, maupun Promo Material. Contoh: biaya retribusi legalitas perizinan display dinas pertanian, donasi/CSR desa penunjang demo plot, atau renovasi mendadak kantor perwakilan area.")

    add_section_header("9. Modul Penggajian Karyawan Lapangan", "kmt/gaji")
    p = doc.add_paragraph()
    p.add_run("Menu khusus bagi ").font.size = Pt(9.5)
    p.add_run("Level 1 (Kepala Departemen / Super Admin)").bold = True
    p.add_run(" untuk mengelola dan merekap payroll bulanan tim lapangan divisi jagung.\n• Menampilkan matriks gaji 12 bulan (Januari s.d. Desember) per karyawan.\n• Memuat informasi status kepegawaian (Aktif/Resign), posisi, dan tanggal join.\n• Terkoneksi langsung membentuk komponen 'Total Biaya Gaji' pada Dashboard Cost / Hasil.\n• Mendukung fitur Import Excel Payroll untuk efisiensi entri data.")

    doc.add_page_break()

    # --- BAB 10 ---
    add_section_header("10. Workflow Verifikasi & Kontrol Keuangan")
    p = doc.add_paragraph()
    p.add_run("Untuk menjamin validitas dan kepatuhan finansial, setiap transaksi DCA dan Biaya Operasional wajib melalui alur persetujuan berikut:")

    flow_table = doc.add_table(rows=5, cols=2)
    flow_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    flow_table.autofit = False

    steps_flow = [
        ("Langkah 1: Input Data oleh ABM", "ABM menginput rincian biaya DCA / Operasional beserta bukti pengeluaran. Status data awal: 'Belum Terverifikasi' (ikon jam kuning)."),
        ("Langkah 2: Review Finance / KADEP", "Admin Keuangan atau KADEP memeriksa kesesuaian kwitansi, nominal biaya, dan saldo uang muka kasbon."),
        ("Langkah 3: Persetujuan Verifikasi", "Jika disetujui, Admin Keuangan mengklik tombol Verifikasi & mengisi Catatan. Status berubah menjadi 'Terverifikasi' (ikon centang hijau)."),
        ("Langkah 4: Penguncian Akses ABM", "Data terkunci secara otomatis. ABM hanya memiliki hak melihat (Read-Only) dan tidak dapat mengedit atau menghapus data."),
        ("Langkah 5: Audit Edit Ulang Otomatis", "Jika Admin Keuangan atau KADEP melakukan koreksi nominal setelah verifikasi, status verifikasi OTOMATIS RESET ke 0 dan wajib diverifikasi ulang demi transparansi audit.")
    ]

    for idx, (s_title, s_desc) in enumerate(steps_flow):
        c1 = flow_table.cell(idx, 0)
        c2 = flow_table.cell(idx, 1)
        c1.width = Inches(2.2)
        c2.width = Inches(4.3)
        set_cell_background(c1, "1E3A8A" if idx == 2 else "F3F4F6")
        set_cell_background(c2, "FFFFFF")
        set_cell_margins(c1, top=60, bottom=60, left=80, right=80)
        set_cell_margins(c2, top=60, bottom=60, left=80, right=80)
        set_cell_border(c1, bottom={"sz": 4, "val": "single", "color": "D1D5DB"},
                            top={"sz": 4, "val": "single", "color": "D1D5DB"},
                            left={"sz": 4, "val": "single", "color": "D1D5DB"},
                            right={"sz": 4, "val": "single", "color": "D1D5DB"})
        set_cell_border(c2, bottom={"sz": 4, "val": "single", "color": "D1D5DB"},
                            top={"sz": 4, "val": "single", "color": "D1D5DB"},
                            left={"sz": 4, "val": "single", "color": "D1D5DB"},
                            right={"sz": 4, "val": "single", "color": "D1D5DB"})

        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_before = Pt(1)
        p1.paragraph_format.space_after = Pt(1)
        r1 = p1.add_run(s_title)
        r1.bold = True
        r1.font.size = Pt(8.5)
        if idx == 2:
            r1.font.color.rgb = RGBColor(255, 255, 255)
        else:
            r1.font.color.rgb = RGBColor(30, 58, 138)

        p2 = c2.paragraphs[0]
        p2.paragraph_format.space_before = Pt(1)
        p2.paragraph_format.space_after = Pt(1)
        r2 = p2.add_run(s_desc)
        r2.font.size = Pt(8.5)
        r2.font.color.rgb = RGBColor(55, 65, 81)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # --- BAB 11 ---
    add_section_header("11. Panduan Import & Export Excel")

    doc.add_paragraph("11.1. Panduan Import Data Massal (Excel)").bold = True
    p = doc.add_paragraph()
    p.add_run("Fitur import tersedia pada modul: ").font.size = Pt(9.5)
    p.add_run("Omset Penjualan, Biaya DCA, Promo Material, Retur, dan Gaji Karyawan.").bold = True

    imp_steps = [
        "Masuk ke halaman modul yang dituju -> klik tombol 'Import Excel'.",
        "Wajib klik tautan 'Download Format Excel (Template)' untuk mendapatkan format kolom resmi sistem.",
        "Isi data pada file template tersebut tanpa mengubah tata letak, susunan, maupun nama judul kolom header.",
        "Simpan file dalam format .xlsx atau .xls.",
        "Upload kembali file melalui form modal import -> klik 'Upload & Proses Data'.",
        "Sistem akan memvalidasi data dan menampilkan notifikasi jumlah baris yang sukses diimpor."
    ]
    for step in imp_steps:
        p_st = doc.add_paragraph(style='List Number')
        p_st.paragraph_format.space_before = Pt(1)
        p_st.paragraph_format.space_after = Pt(2)
        r = p_st.add_run(step)
        r.font.size = Pt(9.5)

    doc.add_paragraph("11.2. Panduan Export Laporan ke Excel").bold = True
    p = doc.add_paragraph()
    p.add_run("Setiap modul tabel dan dashboard dilengkapi tombol ")
    p.add_run("Export Excel").bold = True
    p.add_run(". Data yang diekspor otomatis menyesuaikan filter aktif (Tahun, Periode Bulan, & Wilayah).\nKhusus modul DCA tersedia 3 format ekspor:")
    
    exp_types = [
        ("• Export Ringkasan", "Menghasilkan tabel ringkas daftar pengajuan DCA."),
        ("• Export Rincian", "Menghasilkan dokumen detail item kegiatan per nomor DCA."),
        ("• Export Rekap Berjenjang", "Menghasilkan lembar laporan hierarkis ABM -> MDO -> Kegiatan format resmi siap cetak.")
    ]
    for t1, t2 in exp_types:
        p_ex = doc.add_paragraph()
        p_ex.paragraph_format.left_indent = Inches(0.2)
        p_ex.paragraph_format.space_before = Pt(1)
        p_ex.paragraph_format.space_after = Pt(1)
        p_ex.add_run(t1 + ": ").bold = True
        p_ex.add_run(t2)
        p_ex.runs[0].font.size = Pt(9)
        p_ex.runs[1].font.size = Pt(9)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # --- BAB 12 ---
    add_section_header("12. Pertanyaan Umum (FAQ) & Penanganan Masalah")

    faqs = [
        ("Q1: Mengapa akun saya sebagai ABM tidak dapat melihat menu Omset atau Gaji?",
         "A1: Sesuai kebijakan tata kelola internal PT Karisma Indoagro Universal, data omset nasional dan slip gaji bersifat rahasia (confidential) dan hanya diperuntukkan bagi manajemen pusat (Finance & KADEP). ABM difokuskan pada eksekusi kegiatan promosi (DCA) dan efisiensi operasional wilayahnya."),
        
        ("Q2: Saya salah menginput nominal biaya DCA, tetapi tombol Edit tidak dapat diklik?",
         "A2: Data tersebut telah berstatus 'Terverifikasi' oleh Admin Keuangan sehingga otomatis terkunci. Silakan berkoordinasi dengan Admin Keuangan pusat untuk membuka/membatalkan status verifikasi dokumen tersebut terlebih dahulu."),
        
        ("Q3: Mengapa angka persentase Cost / Hasil di Dashboard berwarna merah?",
         "A3: Indikator persentase otomatis berwarna merah jika rasio pengeluaran biaya terhadap total omset neto melebihi 30.00%. Ini merupakan peringatan sistem bahwa biaya pada periode/wilayah tersebut telah melampaui ambang batas ideal efisiensi."),
         
        ("Q4: Bagaimana cara memastikan nilai Retur memotong omset di Dashboard?",
         "A4: Saat menginput data retur di menu Omset, pastikan memilih opsi 'Kurangi Target: YA'. Dengan opsi ini, omset neto toko terkait otomatis dipotong sebesar nilai retur sehingga angka akumulasi di Dashboard langsung terkoreksi akurat.")
    ]

    for q, a in faqs:
        p_q = doc.add_paragraph()
        p_q.paragraph_format.space_before = Pt(4)
        p_q.paragraph_format.space_after = Pt(1)
        r_q = p_q.add_run(q)
        r_q.bold = True
        r_q.font.size = Pt(9.5)
        r_q.font.color.rgb = RGBColor(30, 58, 138)

        p_a = doc.add_paragraph()
        p_a.paragraph_format.space_before = Pt(1)
        p_a.paragraph_format.space_after = Pt(4)
        p_a.paragraph_format.left_indent = Inches(0.2)
        r_a = p_a.add_run(a)
        r_a.font.size = Pt(9)
        r_a.font.color.rgb = RGBColor(55, 65, 81)

    # --- FOOTER BLOCK ---
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    foot_box = doc.add_table(rows=1, cols=1)
    foot_box.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_f = foot_box.cell(0, 0)
    c_f.width = Inches(6.5)
    set_cell_background(c_f, "F9FAFB")
    set_cell_margins(c_f, top=100, bottom=100, left=150, right=150)
    set_cell_border(c_f, top={"sz": 8, "val": "single", "color": "1E3A8A"},
                         bottom={"val": "none"}, left={"val": "none"}, right={"val": "none"})
    p_f = c_f.paragraphs[0]
    p_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_f.paragraph_format.space_before = Pt(0)
    p_f.paragraph_format.space_after = Pt(0)
    r_f1 = p_f.add_run("Buku Panduan Aplikasi KMT CORN • Karisma ERP\n")
    r_f1.bold = True
    r_f1.font.size = Pt(9)
    r_f1.font.color.rgb = RGBColor(30, 58, 138)
    r_f2 = p_f.add_run("Dikelola & Didistribusikan oleh Tim Pengembang IT — PT Karisma Indoagro Universal")
    r_f2.font.size = Pt(8.5)
    r_f2.font.color.rgb = RGBColor(107, 114, 128)

    # Save documents
    target_path = r"c:\laragon\www\karismaerp\BUKU_PANDUAN_KMT_CORN.docx"
    doc.save(target_path)
    print(f"Document successfully created at {target_path}")

    # Also save as GUIDE_BOOK_KMT_CORN.docx for user convenience
    doc.save(r"c:\laragon\www\karismaerp\GUIDE_BOOK_KMT_CORN.docx")
    print("Also saved as GUIDE_BOOK_KMT_CORN.docx")

if __name__ == "__main__":
    build_word_guide()
